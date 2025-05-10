from fastapi import FastAPI, UploadFile, Form, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from .AI.tailor import enhance_resume
import io
import PyPDF2
from docx import Document
import logging
import traceback
import os
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import time
import tempfile
from typing import Optional

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI()

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # For development only
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Mount static files
app.mount("/static", StaticFiles(directory="frontend"), name="static")

# Serve frontend files
@app.get("/")
async def read_index():
    return FileResponse("frontend/src/pages/index.html")

@app.get("/results")
async def read_results():
    return FileResponse("frontend/src/pages/results.html")

def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        logger.info("Extracting text from PDF...")
        pdf_file = io.BytesIO(file_bytes)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            logger.info(f"PDF Page Content: {page_text[:200]}...")  # Log first 200 chars
            text += page_text + "\n"
        return text.strip()
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=400, detail=f"Error processing PDF file: {str(e)}")

def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        logger.info("Extracting text from DOCX...")
        docx_file = io.BytesIO(file_bytes)
        doc = Document(docx_file)
        text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                logger.info(f"DOCX Paragraph: {paragraph.text[:200]}...")  # Log first 200 chars
                text.append(paragraph.text)
        return "\n".join(text)
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=400, detail=f"Error processing DOCX file: {str(e)}")

def generate_pdf(resume_text: str) -> bytes:
    """Generate a PDF from the resume text."""
    try:
        # Create a buffer to store the PDF
        buffer = io.BytesIO()
        
        # Create the PDF document
        doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
        
        # Create styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=20,
            textColor=colors.HexColor('#2c3e50')
        )
        section_style = ParagraphStyle(
            'CustomSection',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=10,
            textColor=colors.HexColor('#3498db')
        )
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=5,
            leading=14
        )
        
        # Split the resume into sections
        sections = resume_text.split('\n\n')
        story = []
        
        for section in sections:
            if section.startswith('='):
                # Skip separator lines
                continue
            elif section.upper() == section and not section.startswith('['):
                # This is a section header
                story.append(Paragraph(section, section_style))
            else:
                # This is body text
                story.append(Paragraph(section, body_style))
                story.append(Spacer(1, 5))
        
        # Build the PDF
        doc.build(story)
        
        # Get the PDF bytes
        pdf_bytes = buffer.getvalue()
        buffer.close()
        
        return pdf_bytes
    except Exception as e:
        logger.error(f"Error generating PDF: {str(e)}")
        raise

@app.post("/api/enhance-resume")
async def enhance_resume_endpoint(
    resume: UploadFile = File(...),
    job_description: str = Form(...)
):
    try:
        # Validate file type
        if not resume.filename.lower().endswith((".pdf", ".docx")):
            raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported")
        content = await resume.read()
        if resume.filename.lower().endswith('.pdf'):
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            resume_text = ""
            for page in pdf_reader.pages:
                resume_text += page.extract_text() + "\n"
        else:
            docx_file = io.BytesIO(content)
            doc = Document(docx_file)
            resume_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        if not job_description or not job_description.strip():
            raise HTTPException(status_code=400, detail="Job description is required")
        try:
            result = enhance_resume(resume_text, job_description)
            return JSONResponse(content={"analysis": result["analysis"]})
        except Exception as e:
            logger.error(f"Error enhancing resume: {str(e)}")
            logger.error(traceback.format_exc())
            raise HTTPException(status_code=500, detail="Error processing resume")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Unexpected error: {str(e)}")
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail="Internal server error")

@app.get("/api/download-resume")
async def download_resume():
    try:
        # Create a temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
            # Create PDF using the enhanced resume content
            pdf_bytes = create_pdf(enhanced_resume)
            temp_file.write(pdf_bytes)
            temp_file_path = temp_file.name
        
        return FileResponse(
            temp_file_path,
            media_type='application/pdf',
            filename="enhanced_resume.pdf"
        )
    except Exception as e:
        logger.error(f"Error creating PDF: {str(e)}")
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail="Error creating PDF") 